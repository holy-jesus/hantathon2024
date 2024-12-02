from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import RGBColor, Pt, Mm
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn


DOCUMENT_TITLE = (
    "Отчет проверки доступности веб-сайта для пользователей с нарушением зрения"
)
DEFIANCES_PARAGRAPH = "Выявленные нарушения:"
NO_DEFIANCES = "Нарушений нету."
RECOMMENDATIONS_PARAGRAPH = "Рекомендации:"
NO_RECOMMENDATIONS = "Рекомендаций нету."
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(16)


class Report:
    def __init__(self, url: str):
        self.__doc = Document()
        self.__url = url
        self.__defiances = []
        self.__recommendations = []
        self.__screenshot = None

        self.__init_style()
        self.__init_header()

    def __init_style(self):
        style = self.__doc.styles["Normal"]
        style.font.name = FONT_NAME
        style.font.size = FONT_SIZE

    def __init_header(self):
        heading = self.__doc.add_paragraph(DOCUMENT_TITLE)
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        heading.runs[0].font.bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date = datetime.now().strftime("%d.%m.%Y")
        # Страница сайта, дата проверки
        self.__doc.add_paragraph(f"Дата проверки: {date}")
        self.__doc.add_paragraph(f"Страница сайта: {self.__url}")

    def set_screenshot(self, screenshot: bytes) -> None:
        self.__doc.add_paragraph("Скриншот сайта:")
        self.__doc.add_picture(BytesIO(screenshot), width=Mm(self.__get_text_width()))

    def __get_text_width(self):
        """
        Returns the text width in mm.
        """
        section = self.__doc.sections[0]
        return (section.page_width - section.left_margin - section.right_margin) / 36000

    def add_defiance(self, text: str, url: str = None, word: str = None) -> None:
        if not url:
            self.__defiances.append(text)
        else:
            self.__defiances.append({"text": text, "url": url, "word": word})

    def add_recommendation(self, text: str, url: str = None, word: str = None) -> None:
        if not url:
            self.__recommendations.append(text)
        else:
            self.__recommendations.append({"text": text, "url": url, "word": word})

    def render(self) -> BytesIO:
        if self.__defiances:
            v = self.__doc.add_paragraph(DEFIANCES_PARAGRAPH)

            for i, defiance in enumerate(self.__defiances, start=1):
                if isinstance(defiance, dict):
                    text = defiance["text"]
                    word = defiance["word"]
                    url = defiance["url"]
                    paragraph = self.__doc.add_paragraph(f"{i}. {text}")
                    self.__add_hyperlink(paragraph, url, word)
                else:
                    self.__doc.add_paragraph(f"{i}. {defiance}")
        else:
            v = self.__doc.add_paragraph(NO_DEFIANCES)
        v.runs[0].font.bold = True

        if self.__recommendations:
            r = self.__doc.add_paragraph(RECOMMENDATIONS_PARAGRAPH)
            for i, recommendation in enumerate(self.__recommendations, start=1):
                if isinstance(recommendation, dict):
                    word = recommendation["word"]
                    url = recommendation["url"]
                    text = recommendation["text"]
                    paragraph = self.__doc.add_paragraph(f"{i}. {text}")
                    self.__add_hyperlink(paragraph, url, word)
                else:
                    self.__doc.add_paragraph(f"{i}. {recommendation}")
        else:
            r = self.__doc.add_paragraph(NO_RECOMMENDATIONS)
        r.runs[0].font.bold = True

        file = BytesIO()
        self.__doc.save(file)
        file.seek(0)
        return file

    def __add_hyperlink(
        self, paragraph: Paragraph, url: str, text: str, color="0000FF", underline=True
    ):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")

        if color:
            color_element = OxmlElement("w:color")
            color_element.set(qn("w:val"), color)
            r_pr.append(color_element)
        if not underline:
            u_element = OxmlElement("w:u")
            u_element.set(qn("w:val"), "none")
            r_pr.append(u_element)

        run.append(r_pr)
        text_element = OxmlElement("w:t")
        text_element.text = text
        run.append(text_element)
        hyperlink.append(run)

        paragraph._element.append(hyperlink)
