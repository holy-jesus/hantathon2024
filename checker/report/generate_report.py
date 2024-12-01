from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def generate_accessibility_report(link, violations, recommendations):
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(16)
    
    # Заголовок документа
    heading = doc.add_paragraph('Отчет проверки доступности веб-сайта для пользователей с нарушением зрения')
    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0) 
    heading.runs[0].font.bold = True 
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Страница сайта, дата проверки
    doc.add_paragraph(f'Дата проверки: {date}')
    doc.add_paragraph(f'Страница сайта: {link}')
    
    
    # Выявленные нарушения
    v = doc.add_paragraph('Выявленные нарушения:')
    v.runs[0].font.bold = True 
    for i, violation in enumerate(violations, start=1):
        doc.add_paragraph(f'{i}. {violation}')
        # Добавляем картинки 
        # for j in range(1, 4):  # Добавляем 3 картинки
        #     doc.add_paragraph(f'картинка {j}')
    
    # Рекомендации
    r = doc.add_paragraph('Рекомендации:')
    r.runs[0].font.bold = True 
    for i, recommendation in enumerate(recommendations, start=1):
        doc.add_paragraph(f'{i}. {recommendation}')
    
    # Сохраняем документ
    doc.save('report.docx')

# Пример использования функции
link = 'https://example.com'
date = '01.12.2024'
violations = [
    'недостаточная контрастность текста.',
    'отсутствует альтернативный текст для изображений.'
]

recommendations = [
    'увеличить контрастность текста.',
    'добавить альтернативный текст для всех изображений.'
]

generate_accessibility_report(link, violations, recommendations)
